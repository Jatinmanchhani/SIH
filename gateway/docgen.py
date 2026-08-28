"""
docgen.py — generates the actual deliverable: a Word approval note from structured findings.

This is deliberately template-driven, not free-text generation. The model fills in
structured fields (findings, citations, recommendation); this module lays them into a
consistent, professional format. That split matters for two reasons:
  1. It's what makes the output look like a real approval note instead of a chat transcript.
  2. It's the seam where "human-in-the-loop" actually lives — every finding below carries
     its source citation, and the note is explicitly marked DRAFT until a person signs off.
"""

from __future__ import annotations
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


@dataclass
class Finding:
    text: str
    source: str  # e.g. "Scanned inspection report, page 2" or "SOP-MECH-014"


@dataclass
class ApprovalNote:
    title: str
    reference_no: str
    prepared_for: str
    findings: list[Finding]
    recommendation: str
    prepared_by_system: str = "AI Workbench (draft — not yet reviewed)"
    date_str: str = field(default_factory=lambda: date.today().isoformat())


def _set_run(run, size=11, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def build_approval_note(note: ApprovalNote, output_path: Path) -> Path:
    doc = Document()

    # US Letter, per docx conventions — avoids the default-A4 surprise.
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    # Header block
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(note.title)
    _set_run(run, size=16, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Reference: {note.reference_no}    |    Date: {note.date_str}")
    _set_run(run, size=10, color=(90, 90, 90))

    # DRAFT banner — this is not decorative, it's the human-in-the-loop control.
    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = banner.add_run(
        "DRAFT — machine-generated. Requires reviewer sign-off before this note is valid."
    )
    _set_run(run, size=10, bold=True, color=(180, 40, 40))

    doc.add_paragraph()  # spacer

    p = doc.add_paragraph()
    p.add_run("Prepared for: ").bold = True
    p.add_run(note.prepared_for)

    p = doc.add_paragraph()
    p.add_run("Prepared by: ").bold = True
    p.add_run(note.prepared_by_system)

    # Findings section
    doc.add_heading("Key Findings", level=2)
    for i, f in enumerate(note.findings, start=1):
        para = doc.add_paragraph(style="List Number")
        run = para.add_run(f.text)
        _set_run(run, size=11)
        cite = doc.add_paragraph()
        cite.paragraph_format.left_indent = Inches(0.35)
        run = cite.add_run(f"Source: {f.source}")
        _set_run(run, size=9, color=(110, 110, 110))

    # Recommendation section
    doc.add_heading("Recommendation", level=2)
    doc.add_paragraph(note.recommendation)

    # Sign-off block — the actual human-in-the-loop gate
    doc.add_heading("Reviewer Sign-off", level=2)
    table = doc.add_table(rows=2, cols=2)
    table.style = "Light Grid Accent 1"
    table.rows[0].cells[0].text = "Reviewed by (name)"
    table.rows[0].cells[1].text = "Decision (Approve / Reject / Return for revision)"
    table.rows[1].cells[0].text = ""
    table.rows[1].cells[1].text = ""

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


if __name__ == "__main__":
    # Self-test: build a realistic sample approval note end to end.
    sample = ApprovalNote(
        title="Approval Note — Vendor Certificate Exception, PO 4471228",
        reference_no="AN-2026-0341",
        prepared_for="Category Manager, Mechanical Spares",
        findings=[
            Finding(
                text="Vendor's ISO 9001 certificate expired 2026-08-11, 14 days prior to PO release date.",
                source="Scanned vendor certificate, uploaded 2026-08-25",
            ),
            Finding(
                text="No renewed certificate received within the standard 10 working day response window.",
                source="Vendor correspondence log, entries dated 2026-08-12 to 2026-08-22",
            ),
            Finding(
                text="Per SOP-PROC-007 section 3(d), proceeding requires written Procurement Head sign-off, "
                     "valid for this PO only.",
                source="SOP-PROC-007: Vendor Certificate Validity and Onboarding Exceptions",
            ),
        ],
        recommendation=(
            "Recommend a single-PO exception under SOP-PROC-007 3(d), contingent on Procurement Head "
            "sign-off below. This exception does not waive certificate renewal for future POs — "
            "vendor record remains on-hold pending a valid ISO 9001 certificate."
        ),
    )
    out = build_approval_note(sample, Path(__file__).parent.parent.parent / "sample_data" / "sample_approval_note.docx")
    print(f"Wrote {out}")
